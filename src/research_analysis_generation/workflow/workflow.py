import os
import sys
from datetime import datetime
from typing import Optional
from langgraph.types import Send
from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain_core.messages import get_buffer_string
from langchain_community.tools.tavily_search import TavilySearchResults
import re

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from research_analysis_generation.exception.custom_exception import ResearchAnalysisException
from research_analysis_generation.workflow.interview_workflow import InterviewGraphBuilder
from research_analysis_generation.schema.model import (
    Perspectives,
    GenerateAnalystsState,
    ResearchGraphState
)
from research_analysis_generation.utils.api_key_manager import ApiKeyManager
from research_analysis_generation.logger import GLOBAL_LOGGER
from research_analysis_generation.utils.model_loader import ModelLoader
from research_analysis_generation.prompt_lib.prompt import (
    CREATE_ANALYSTS_PROMPT,
    REPORT_WRITER_INSTRUCTIONS,
    INTRO_CONCLUSION_INSTRUCTIONS
)

current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.abspath(os.path.join(current_dir, "../../"))
sys.path.append(project_root)
    
class ReportGenerator:
    
    def __init__(self, llm):
        self.llm = llm
        self.memory = MemorySaver()
        self.logger = GLOBAL_LOGGER.bind(module = "ReportGenerator")
        
        self.tavily_search = TavilySearchResults(api_key=ApiKeyManager().get("TAVILY_API_KEY"))
    
    def create_analyst(self, state: GenerateAnalystsState):
        topic = state["topic"]
        max_analysts = state["max_analysts"]
        human_analyst_feedback = state.get("human_analyst_feedback", "")
        try:
            self.logger.info(f"Creating analysts for topic: {topic} with max analysts: {max_analysts}")
            
            structured_llm = self.llm.with_structured_output(Perspectives)
            system_message = CREATE_ANALYSTS_PROMPT.render(
                topic= topic,
                human_analyst_feedback= human_analyst_feedback,
                max_analysts= max_analysts
            )
            analysts = structured_llm.invoke([
                SystemMessage(content= system_message),
                HumanMessage(content= "create a list of analysts")
            ])
            return {
                "analysts": analysts.analysts
                }
        except Exception as e:
            self.logger.error(f"Error creating analysts: {e}")
            raise ResearchAnalysisException("Failed to create analysts", e)
        
    def human_feedback(self):
        """This node is a placeholder for human feedback. It does not perform any operations itself, but serves as an interruption point in the graph where human feedback can be collected and added to the state before proceeding with the report generation process.
        """
        try:
            self.logger.info("Reached human feedback node. Awaiting human feedback before proceeding.")
        except Exception as e:
            self.logger.error(f"Error in human feedback node: {e}")
            raise ResearchAnalysisException("Failed at human feedback node", e)
    
    def write_introduction(self, state: ResearchGraphState):
        """
        Writes the introduction section of the report. 
        This is a placeholder function that can be expanded in the future to generate an introduction based on the topic and sections of the report. Currently, it does not perform any operations.
        """
        try:
            sections = state.get("sections", [])
            topic = state.get("topic", [])
            formatted_sections = "\n\n".join([f"{s}" for s in sections])
            
            self.logger.info(f"Writing introduction for topic: {topic} with sections: {sections}")  
            
            system_prompt = INTRO_CONCLUSION_INSTRUCTIONS.render(
                topic=topic,
                formatted_sections= formatted_sections)
            
            introduction = self.llm.invoke([
                SystemMessage(content= system_prompt),
                HumanMessage(content="Write the introduction section for the report based on the provided sections.")])
            
            self.logger.info("Introduction successfully written.")
            return {"introduction": introduction.content}
        except Exception as e:
            self.logger.error(f"Error writing introduction: {e}")
            raise ResearchAnalysisException("Failed to write introduction", e)
    
    def write_conclusion(self, state: ResearchGraphState):
        """ Writes the conclusion section of the report."""
        try:
            sections = state["sections"]
            topic = state["topic"]
            formatted_sections = "\n\n".join([f"{s}" for s in sections])
            
            self.logger.info(f"Writing conclusion for topic: {topic} with sections: {sections}")
            
            system_prompt = INTRO_CONCLUSION_INSTRUCTIONS.render(
                topic=topic,
                formatted_sections= formatted_sections)
            
            conclustion = self.llm.invoke([
                SystemMessage(content= system_prompt),
                HumanMessage(content="Write the conclusion section for the report based on the provided sections.")])       
            
            self.logger.info("Conclusion successfully written.")
            return {"conclusion": conclustion.content}
        except Exception as e:
            self.logger.error(f"Error writing conclusion: {e}")
            raise ResearchAnalysisException("Failed to write conclusion", e)
    
    def write_report(self, state: ResearchGraphState):
        sections = state.get("sections", [])
        topic = state.get("topic", "")
        
        try:
            if not sections:
                sections_text = "No sections available."
                self.logger.warning("No sections found in state for report writing.")
            else:
                sections_text = "\n\n".join(
                    [
                        f"## {section.title}\n\n{section.content}"
                        for section in sections
                    ]
                )
            self.logger.info(f"Writing report for topic: {topic} with sections: {sections}")
            
            system_prompt = REPORT_WRITER_INSTRUCTIONS.render(
                topic=topic
                )
            self.logger.debug(f"System prompt for report writing: {system_prompt}")
            
            report = self.llm.invoke([
                SystemMessage(content= system_prompt),
                HumanMessage(content=sections_text)
            ])
            
            self.logger.info("Report successfully written.")
            return {"report": report.content}
        
        except Exception as e:
            self.logger.error(f"Error writing report: {e}")
            raise ResearchAnalysisException("Failed to write report", e)
    
    def finalize_report(self, state: ResearchGraphState):
        """
        Finalizes the report by assembling all the sections, introduction, content and conclusion into a cohesive final report. 
        """
        try:
            self.logger.info("Finalizing report by assembling all sections, introduction, and conclusion.")
            content = state.get("content", "")
            
            if content.startswith("## Insights"):
                content = content.replace("## Insights", "").strip()
                
            sources = None
            
            if "## Sources" in content:
                try:
                    content, sources = content.split("\n## Sources\n")
                except Exception as e:
                    self.logger.warning(f"Failed to extract sources from content: {e}")
                    sources = None
                    
            final_report = (
                state["introduction"] + "\n\n" +
                content + "\n\n" +
                state["conclusion"]
            )
            
            if sources:
                final_report += "\n\n## Sources\n" + sources
                
            self.logger.info("Report successfully finalized.")
            return {"final_report": final_report}
        except Exception as e:
            self.logger.error(f"Error finalizing report: {e}")
            raise ResearchAnalysisException("Failed to finalize report", e)
    
    def save_report(self, final_report: str, topic: str, format: str = "docx"):
        """Save the report as DOCX or PDF, each in its own subfolder."""
        try:
            self.logger.info("Saving report", topic=topic, format=format)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_topic = re.sub(r'[\\/*?:"<>|]', "_", topic)
            base_name = f"{safe_topic.replace(' ', '_')}_{timestamp}"

            project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../"))
            root_dir = os.path.join(project_root, "generated_report")

            report_folder = os.path.join(root_dir, base_name)
            os.makedirs(report_folder, exist_ok=True)

            file_path = os.path.join(report_folder, f"{base_name}.{format}")

            if format == "docx":
                self.__save_as_docx(final_report, file_path)
            elif format == "pdf":
                self.__save_as_pdf(final_report, file_path)
            else:
                raise ValueError("Invalid format. Use 'docx' or 'pdf'.")

            self.logger.info("Report saved successfully", path=file_path)
            return file_path

        except Exception as e:
            self.logger.error("Error saving report", error=str(e))
            raise ResearchAnalysisException("Failed to save report file", e)

    def __save_as_docx(self, text: str, file_path: str):
        """Save as DOCX."""
        try:
            doc = Document()
            for line in text.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                else:
                    doc.add_paragraph(line)
            doc.save(file_path)
        except Exception as e:
            self.logger.error("DOCX save failed", path=file_path, error=str(e))
            raise ResearchAnalysisException("Error saving DOCX report", e)

    def __save_as_pdf(self, text: str, file_path: str):
        """Save as PDF with centered text block, wrapping, and clean layout."""
        from textwrap import wrap
        try:
            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter

            left_margin = 80
            right_margin = 80
            usable_width = width - left_margin - right_margin
            top_margin = 70
            bottom_margin = 60
            y = height - top_margin

            normal_font = "Helvetica"
            bold_font = "Helvetica-Bold"
            line_height = 15

            lines = text.split("\n")
            for raw_line in lines:
                line = raw_line.strip()
                if not line:
                    y -= line_height
                    continue

                if line.startswith("# "):
                    font = bold_font
                    size = 16
                    line = line[2:].strip()
                elif line.startswith("## "):
                    font = bold_font
                    size = 13
                    line = line[3:].strip()
                else:
                    font = normal_font
                    size = 11

                c.setFont(font, size)
                wrapped_lines = wrap(line, width=int(usable_width / (size * 0.55)))

                for wline in wrapped_lines:
                    if y < bottom_margin:
                        c.showPage()
                        c.setFont(font, size)
                        y = height - top_margin

                    text_width = c.stringWidth(wline, font, size)
                    x = (width - text_width) / 2 

                    c.drawString(x, y, wline)
                    y -= line_height

            for page_num in range(1, c.getPageNumber() + 1):
                c.setFont("Helvetica", 9)
                c.drawCentredString(width / 2, 25, f"Page {page_num}")

            c.save()
            self.logger.info("Centered PDF saved successfully", path=file_path)

        except Exception as e:
            self.logger.error("PDF save failed", path=file_path, error=str(e))
            raise ResearchAnalysisException("Error saving PDF report", e)
    
    def initiate_all_interviews(self, state: ResearchGraphState):
        topic = state.get("topic", "Untitled Topic")
        analysts = state.get("analysts", [])

        if not analysts:
            self.logger.warning("No analysts found — skipping interviews")
            return END

        return [
            Send(
                "conduct_interview",
                {
                    "analyst": analyst,
                    "messages": [
                        HumanMessage(
                            content=f"So, let's discuss about {topic}."
                        )
                    ],
                    "max_num_turns": 2,
                    "context": [],
                    "interview": "",
                    "sections": [],
                },
            )
            for analyst in analysts
        ]
    def build_graph(self):
        """ Construct the report workflow graph for the research analysis report generation process. """
        try:
            self.logger.info("Starting to build the report workflow graph.")
            interview_graph = InterviewGraphBuilder(self.llm, self.tavily_search).build_graph()

            graph = StateGraph(ResearchGraphState)
            
            graph.add_node("create_analyst", self.create_analyst)
            graph.add_node("human_feedback", self.human_feedback)
            graph.add_node("conduct_interview", interview_graph)
            graph.add_node("write_report", self.write_report)
            graph.add_node("write_introduction", self.write_introduction)
            graph.add_node("write_conclusion", self.write_conclusion)
            graph.add_node("finalize_report", self.finalize_report)
            
            graph.add_edge(START, "create_analyst")
            graph.add_edge("create_analyst", "human_feedback")
            graph.add_conditional_edges("human_feedback", 
                                        self.initiate_all_interviews,
                                        ["conduct_interview", END])
            graph.add_edge("conduct_interview", "write_report")
            graph.add_edge("conduct_interview", "write_introduction")
            graph.add_edge("conduct_interview", "write_conclusion")
            graph.add_edge(["write_report","write_introduction", "write_conclusion"], "finalize_report")
            graph.add_edge("finalize_report", END)
            
            report_graph = graph.compile(interrupt_before=["human_feedback"], checkpointer= self.memory)
            self.logger.info("Report workflow graph successfully built and compiled.")
            
            return report_graph
        except Exception as e:
            self.logger.error(f"Error building report workflow graph: {e}")
            raise ResearchAnalysisException("Failed to build report workflow graph", e)
    
    
if __name__ == "__main__":
    try:
        llm = ModelLoader().load_llm()
        reporter = ReportGenerator(llm)
        graph = reporter.build_graph()

        topic = "Impact of LLMs over the Future of Jobs?"
        thread = {"configurable": {"thread_id": "1"}}
        reporter.logger.info("Starting report generation pipeline", topic=topic)

        for _ in graph.stream({"topic": topic, "max_analysts": 3}, thread, stream_mode="values"):
            pass

        state = graph.get_state(thread)
        feedback = input("\n Enter your feedback or press Enter to continue: ").strip()
        graph.update_state(thread, {"human_analyst_feedback": feedback}, as_node="human_feedback")

        for _ in graph.stream(None, thread, stream_mode="values"):
            pass

        final_state = graph.get_state(thread)
        final_report = final_state.values.get("final_report")

        if final_report:
            reporter.logger.info("Report generated successfully")
            reporter.save_report(final_report, topic, "docx")
            reporter.save_report(final_report, topic, "pdf")
        else:
            reporter.logger.error("No report content generated")

    except Exception as e:
        GLOBAL_LOGGER.error("Fatal error in main execution", error=str(e))
        raise ResearchAnalysisException("Autonomous report generation pipeline failed", e)
    